"""Generate sample RFP files for testing Stage 3.

Run from the backend folder:
    .venv\\Scripts\\python.exe sample_data\\make_rfp_samples.py

Produces:
    sample_data/rfp_sample.xlsx     (English column headers)
    sample_data/rfp_sample_ar.docx  (Arabic table headers)
Scope lines are written so Stage 4 can later decompose them into assemblies.
"""

import os

from docx import Document
from openpyxl import Workbook

HERE = os.path.dirname(__file__)

# (description, quantity, unit) — mixed EN/AR, assembly-style scope lines.
SCOPE = [
    ("Supply and install 24-way distribution board complete with main breaker", 2, "Each"),
    ("Supply and install copper power cable 3-core 2.5mm2 including conduit", 150, "m"),
    ("توريد وتركيب مقابس كهربائية مزدوجة", 40, "Each"),
    ("Paint interior walls with white emulsion, two coats", 320, "m2"),
    ("توريد وتركيب بلاط بورسلان للأرضيات مقاس 60×60 سم", 85, "m2"),
    ("Supply and lay PVC drainage pipe 110mm including fittings", 60, "m"),
]


def make_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "Scope of Work"
    ws.append(["Item Description", "Quantity", "Unit"])
    for desc, qty, unit in SCOPE:
        ws.append([desc, qty, unit])
    path = os.path.join(HERE, "rfp_sample.xlsx")
    wb.save(path)
    print("wrote", path)


def make_docx():
    doc = Document()
    doc.add_heading("Request for Proposal — Scope of Work", level=1)
    doc.add_paragraph("The contractor shall supply and install the following:")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "الوصف", "الكمية", "الوحدة"
    for desc, qty, unit in SCOPE:
        cells = table.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = desc, str(qty), unit
    path = os.path.join(HERE, "rfp_sample_ar.docx")
    doc.save(path)
    print("wrote", path)


if __name__ == "__main__":
    make_xlsx()
    make_docx()
