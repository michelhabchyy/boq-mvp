"""Parse an uploaded RFP / scope-of-work into structured scope lines.

Supports digital .xlsx and .docx (no OCR — scanned PDFs come later). The goal
is a clean list of {description, quantity, unit}, with Arabic/English preserved.

Strategy:
  * Tables (Excel sheets, Word tables): find a header row using bilingual
    aliases, map the description/quantity/unit columns, then read data rows.
    If no header is recognisable, fall back to positional columns.
  * Word with no usable table: treat list/numbered paragraphs as descriptions.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field

from docx import Document
from openpyxl import load_workbook

# Canonical column -> accepted header spellings (EN + AR). EN matched after
# normalisation (lowercased, spaces/hyphens -> underscore); AR matched as-is.
HEADER_ALIASES: dict[str, set[str]] = {
    "description": {
        "description", "desc", "scope", "scope_of_work", "item", "item_description",
        "works", "work", "particulars", "bill_description", "boq_description",
        "الوصف", "وصف", "البند", "بيان", "الاعمال", "الأعمال", "وصف_الاعمال",
    },
    "quantity": {
        "quantity", "qty", "qnty", "no", "nos", "count",
        "الكمية", "كمية", "العدد", "عدد",
    },
    "unit": {
        "unit", "uom", "units", "u_m",
        "الوحدة", "وحدة", "الوحده",
    },
}

_MAX_HEADER_SCAN = 15  # rows to scan looking for a header


@dataclass
class ScopeLine:
    description: str
    quantity: float | None = None
    unit: str | None = None


@dataclass
class RFPParseResult:
    source_type: str  # "xlsx" or "docx"
    scope_lines: list[ScopeLine] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


def _normalize_header(h: str) -> str:
    h = (h or "").strip().lower()
    for ch in (" ", "-", "/", "."):
        h = h.replace(ch, "_")
    while "__" in h:
        h = h.replace("__", "_")
    return h.strip("_")


# Build one lookup: every alias (normalised) -> canonical column name.
# Arabic has no case/space normalisation issues, so the same normaliser is safe.
_ALIAS_LOOKUP = {
    _normalize_header(alias): canon
    for canon, aliases in HEADER_ALIASES.items()
    for alias in aliases
}


def _match_header_cell(cell: str) -> str | None:
    return _ALIAS_LOOKUP.get(_normalize_header(cell))


def _detect_columns(row: list[str]) -> dict[str, int]:
    """Map canonical column -> index for a candidate header row."""
    col_map: dict[str, int] = {}
    for idx, cell in enumerate(row):
        canon = _match_header_cell(cell)
        if canon and canon not in col_map:
            col_map[canon] = idx
    return col_map


def _parse_quantity(value: str) -> float | None:
    if value is None:
        return None
    # grab the first number in the cell (handles "150", "150 m", "1,250.5")
    m = re.search(r"-?\d[\d,]*\.?\d*", str(value))
    if not m:
        return None
    try:
        return float(m.group(0).replace(",", ""))
    except ValueError:
        return None


def _lines_from_table(rows: list[list[str]]) -> tuple[list[ScopeLine], str | None]:
    """Extract scope lines from a 2D table. Returns (lines, warning|None)."""
    rows = [r for r in rows if any((c or "").strip() for c in r)]
    if not rows:
        return [], None

    # Look for a header row within the first few rows.
    header_idx, col_map = None, {}
    for i, row in enumerate(rows[:_MAX_HEADER_SCAN]):
        cm = _detect_columns(row)
        if "description" in cm:
            header_idx, col_map = i, cm
            break

    warning = None
    if header_idx is None:
        # No recognisable header: assume positional [description, quantity, unit].
        col_map = {"description": 0}
        if len(rows[0]) > 1:
            col_map["quantity"] = 1
        if len(rows[0]) > 2:
            col_map["unit"] = 2
        data_rows = rows
        warning = (
            "No recognisable header found; assumed columns by position "
            "(1st=description, 2nd=quantity, 3rd=unit)."
        )
    else:
        data_rows = rows[header_idx + 1 :]

    d_i = col_map["description"]
    q_i = col_map.get("quantity")
    u_i = col_map.get("unit")

    lines: list[ScopeLine] = []
    for row in data_rows:
        desc = (row[d_i].strip() if d_i < len(row) else "")
        if not desc:
            continue
        qty = _parse_quantity(row[q_i]) if q_i is not None and q_i < len(row) else None
        unit = (row[u_i].strip() or None) if u_i is not None and u_i < len(row) else None
        lines.append(ScopeLine(description=desc, quantity=qty, unit=unit))
    return lines, warning


def _read_xlsx(content: bytes) -> RFPParseResult:
    wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    ws = wb.active
    sheet_title, sheet_count = ws.title, len(wb.sheetnames)
    rows = [
        ["" if c is None else str(c).strip() for c in r]
        for r in ws.iter_rows(values_only=True)
    ]
    wb.close()

    result = RFPParseResult(source_type="xlsx")
    lines, warning = _lines_from_table(rows)
    result.scope_lines = lines
    if warning:
        result.warnings.append(f"[sheet '{sheet_title}'] {warning}")
    if sheet_count > 1:
        result.warnings.append(
            f"Workbook has {sheet_count} sheets; only the active sheet "
            f"'{sheet_title}' was parsed."
        )
    return result


def _read_docx(content: bytes) -> RFPParseResult:
    doc = Document(io.BytesIO(content))
    result = RFPParseResult(source_type="docx")

    # 1) Tables first.
    for t_idx, table in enumerate(doc.tables):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        lines, warning = _lines_from_table(rows)
        result.scope_lines.extend(lines)
        if warning:
            result.warnings.append(f"[table {t_idx + 1}] {warning}")

    # 2) If tables yielded nothing, fall back to list-style paragraphs.
    if not result.scope_lines:
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            style = (p.style.name or "").lower() if p.style else ""
            is_listish = "list" in style or re.match(r"^\s*(\d+[.)]|[-•*])\s+", text)
            if is_listish or len(text) > 3:
                result.scope_lines.append(ScopeLine(description=text))
        if result.scope_lines:
            result.warnings.append(
                "No tables found; extracted paragraphs as descriptions "
                "(quantity/unit not detected — review these lines)."
            )

    return result


def parse_rfp_file(filename: str, content: bytes) -> RFPParseResult:
    name = (filename or "").lower()
    if name.endswith(".xlsx") or name.endswith(".xlsm"):
        return _read_xlsx(content)
    if name.endswith(".docx"):
        return _read_docx(content)
    raise ValueError("Unsupported file type. Upload a .xlsx or .docx file.")


def extract_full_text(filename: str, content: bytes) -> str:
    """Flatten a document to plain text (paragraphs + tables) for AI analysis."""
    name = (filename or "").lower()
    parts: list[str] = []
    if name.endswith((".xlsx", ".xlsm")):
        wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        for ws in wb.worksheets:
            parts.append(f"# Sheet: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    parts.append(" | ".join(cells))
        wb.close()
    elif name.endswith(".docx"):
        doc = Document(io.BytesIO(content))
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                parts.append(t)
        for i, table in enumerate(doc.tables, start=1):
            parts.append(f"[Table {i}]")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
    else:
        raise ValueError("Unsupported file type. Upload a .xlsx or .docx file.")
    return "\n".join(parts)
